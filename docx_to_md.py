import os
import re
from docx import Document
from docx.shared import Inches
import shutil
from PIL import Image
import base64
import io
import glob
import sys

def clean_filename(filename):
    """清理文件名，移除不合法字符"""
    return re.sub(r'[\\/:*?"<>|]', '_', filename)

def extract_images(doc, output_folder):
    """从docx文件中提取图片并保存到指定文件夹"""
    image_paths = []
    
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_filename = f"image_{len(image_paths) + 1}.png"
            image_path = os.path.join(output_folder, image_filename)
            
            with open(image_path, "wb") as f:
                f.write(image_data)
            
            image_paths.append((rel.rId, image_path))
    
    return image_paths

def docx_to_markdown(docx_path, markdown_path, images_folder):
    """将docx文件转换为markdown并处理图像"""
    doc = Document(docx_path)
    
    # 提取图片
    image_paths = extract_images(doc, images_folder)
    image_dict = {rId: path for rId, path in image_paths}
    
    # 创建markdown内容
    markdown_content = []
    
    # 处理文档标题
    markdown_content.append(f"# {os.path.splitext(os.path.basename(docx_path))[0]}\n\n")
    
    # 处理段落
    image_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检查段落样式，处理标题
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            level = 1  # 默认级别
            if para.style.name[-1].isdigit():
                level = int(para.style.name[-1])
            markdown_content.append(f"{'#' * level} {text}\n\n")
        else:
            # 常规段落
            markdown_content.append(f"{text}\n\n")
            
        # 检查图片
        for run in para.runs:
            if run.element.xpath('.//a:blip'):
                for blip in run.element.xpath('.//a:blip'):
                    rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rId and rId in image_dict:
                        image_count += 1
                        image_path = os.path.relpath(image_dict[rId], os.path.dirname(markdown_path)).replace("\\", "/")
                        markdown_content.append(f"![图片{image_count}]({image_path})\n\n")
    
    # 处理表格
    for table in doc.tables:
        table_md = []
        # 表头
        header_row = []
        for cell in table.rows[0].cells:
            header_row.append(cell.text.strip() or " ")
        table_md.append("| " + " | ".join(header_row) + " |")
        
        # 分隔行
        table_md.append("| " + " | ".join(["---"] * len(header_row)) + " |")
        
        # 表格内容
        for row in table.rows[1:]:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip() or " ")
            table_md.append("| " + " | ".join(row_data) + " |")
        
        markdown_content.append("\n".join(table_md) + "\n\n")
    
    # 写入markdown文件
    with open(markdown_path, "w", encoding="utf-8") as f:
        f.write("".join(markdown_content))
    
    print(f"已将 {docx_path} 转换为 {markdown_path}")
    print(f"提取了 {len(image_paths)} 张图片到 {images_folder}")

def main():
    # 文件夹路径
    documents_folder = "documents"
    images_folder = os.path.join(documents_folder, "images")
    
    # 确保images文件夹存在
    if not os.path.exists(images_folder):
        os.makedirs(images_folder)
    
    # 查找所有docx文件
    docx_files = glob.glob(os.path.join(documents_folder, "*.docx"))
    
    # 处理每个docx文件
    for docx_path in docx_files:
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        markdown_path = os.path.join(documents_folder, f"{base_name}.md")
        
        # 为每个文档创建单独的图片文件夹
        doc_images_folder = os.path.join(images_folder, base_name)
        
        print(f"正在处理: {docx_path}")
        docx_to_markdown(docx_path, markdown_path, doc_images_folder)

if __name__ == "__main__":
    main() 